"""
FX281 Podcast Processing Backend v13.0
- Local FunASR (SenseVoice) STT + DeepSeek analysis + Word/MP3 export
- Three-level deletion suggestion (keep/mild/strong)
- Task persistence to JSON file
- History API
"""

import os
import sys
import io
import json
import re
import uuid
import tempfile
import traceback
import logging
import asyncio
import platform
from typing import List, Dict, Optional
from datetime import datetime
import httpx

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger("fx281")

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY", "")

LLM_PROVIDER = "deepseek"
LLM_MODEL = os.getenv("LLM_MODEL", "deepseek-v4-pro")
ASR_MODE = os.getenv("ASR_MODE", "local")

if not DEEPSEEK_API_KEY:
    logger.warning("DEEPSEEK_API_KEY 未设置，云端文本分析将不可用")
if ASR_MODE == "dashscope" and not DASHSCOPE_API_KEY:
    logger.warning("ASR_MODE=dashscope 但 DASHSCOPE_API_KEY 未设置，云端转写将失败")

DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
os.makedirs(DATA_DIR, exist_ok=True)
HISTORY_FILE = os.path.join(DATA_DIR, "history.json")

os.environ["MODELSCOPE_CACHE"] = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "modelscope_cache")

from openai import OpenAI
deepseek_client = (
    OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    if DEEPSEEK_API_KEY else None
)

# Make ffmpeg/ffprobe discoverable in PATH (used by funasr's _load_audio_ffmpeg and pydub's mediainfo)
import shutil as _shutil
# 1. Add Homebrew bin (macOS Apple Silicon) to PATH if present
_brew_bin = "/opt/homebrew/bin"
if os.path.isdir(_brew_bin) and _brew_bin not in os.environ.get("PATH", ""):
    os.environ["PATH"] = _brew_bin + os.pathsep + os.environ.get("PATH", "")
# 2. Fall back to bundled imageio-ffmpeg if system ffmpeg still not found
if not _shutil.which("ffmpeg"):
    try:
        import imageio_ffmpeg
        _ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        _ffmpeg_dir = os.path.dirname(_ffmpeg_exe)
        _ffmpeg_link = os.path.join(_ffmpeg_dir, "ffmpeg")
        if not os.path.exists(_ffmpeg_link):
            os.symlink(os.path.basename(_ffmpeg_exe), _ffmpeg_link)
        os.environ["PATH"] = _ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
        logger.info(f"System ffmpeg not found; added bundled to PATH: {_ffmpeg_dir}")
    except Exception as _e:
        logger.warning(f"Could not configure bundled ffmpeg: {_e}")
else:
    logger.info(f"Using system ffmpeg: {_shutil.which('ffmpeg')}, ffprobe: {_shutil.which('ffprobe')}")

_local_asr_model = None
_local_vad_model = None
_local_asr_failed = False

def _get_local_asr():
    global _local_asr_model, _local_asr_failed
    if _local_asr_model is not None:
        return _local_asr_model
    if _local_asr_failed:
        return None
    try:
        from funasr import AutoModel
        logger.info("Loading local SenseVoice model...")
        _local_asr_model = AutoModel(
            model="iic/SenseVoiceSmall",
            device="cuda" if _check_cuda() else "cpu",
        )
        logger.info("Local SenseVoice model loaded successfully!")
        return _local_asr_model
    except Exception as e:
        logger.warning(f"Failed to load local ASR model: {e}")
        _local_asr_failed = True
        return None

def _get_local_vad():
    global _local_vad_model
    if _local_vad_model is not None:
        return _local_vad_model
    try:
        from funasr import AutoModel
        logger.info("Loading local VAD model...")
        _local_vad_model = AutoModel(
            model="iic/speech_fsmn_vad_zh-cn-16k-common-pytorch",
            device="cuda" if _check_cuda() else "cpu",
        )
        logger.info("Local VAD model loaded successfully!")
        return _local_vad_model
    except Exception as e:
        logger.warning(f"Failed to load VAD model: {e}")
        _local_vad_model = False
        return None

def _check_cuda():
    try:
        import torch
        return torch.cuda.is_available()
    except:
        return False

tasks: Dict[str, dict] = {}

app = FastAPI(title="FX281 API", version="13.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

DEEPSEEK_BATCH_PROMPT = """你是专业音频剪辑师和内容策划专家，负责播客与访谈节目的粗剪。逐句分析以下转录，标注说话人和剪辑建议。

{transcript_text}

## 说话人
根据对话推断。host=主持人(提问/引导/总结/捧哏)，guest=嘉宾(回答/讲述/分享)。用 Speaker_A/B/C 标识。

## 剪辑建议
**strong（强烈删减，必须剪掉）**
- 准备测试：喂喂喂、能听到吗、调麦克风、开始录了吗
- 意外打断：咳嗽喝水、外人闯入、外卖到了、接个电话
- 偏离主题：与核心话题完全无关的闲聊
- 风险敏感：脏话谩骂、政治敏感、得罪人的不当言论

**mild（一般删减，后期优化）**
- 严重口癖：大量无意义的"然后""就是说""那个""额""啊"，影响连贯
- 结巴重复：自我重复（"我我觉得，我觉得..."），保留完整那句即可
- 逻辑断层：说到一半转话题，前半句无实质信息

**keep（保留，核心骨架）**
- 核心观点：表达清晰、逻辑完整的关键信息
- 自然互动：推动话题的提问、恰当捧哏（"确实""我同意"）
- 情绪价值：活跃气氛、展现性格的幽默或感叹

原因标签：filler(口癖) echo(附和) noise(杂音) redundant(冗余) off_topic(离题) stutter(结巴) sensitive(敏感)
mild/strong 需填 reason 和 reasonDetail(10字内解读)。

返回 JSON：
{{"segments":[{{"id":1,"speaker":"Speaker_A","speakerRole":"host","suggestion":"keep","reason":null,"reasonDetail":null}}]}}
只输出 JSON。"""

DEEPSEEK_CHAPTER_PROMPT = """你是播客编辑。根据以下完整转录，划分5-7个章节。

{transcript_text}

要求：
- 按话题自然转折划分，不按时间均匀切分
- 标题5-10字，具体不泛泛（如"早年求学经历"而非"第一部分"）
- startTime/endTime 取该章节首尾句时间，连续不遗漏
- 即使全篇一个主题，也从不同角度切出至少5章

返回 JSON：
{{"chapters":[{{"title":"开场介绍","startTime":0,"endTime":120}}]}}
只输出 JSON。"""


DEEPSEEK_SUMMARY_PROMPT = """你是播客编辑。根据以下转录，写一段150-250字的内容概述。

{transcript_text}

要求：概括主题、人物、关键话题、整体氛围。简洁流畅，不罗列时间戳。

返回 JSON：
{{"summary":"这是一段关于...的访谈..."}}
只输出 JSON。"""


def format_time(seconds: float) -> str:
    if seconds >= 3600:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours}:{minutes:02d}:{secs:02d}"
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes}:{secs:02d}"


def merge_segments(raw: List[dict], min_chars: int = 15, max_gap: float = 2.0, max_chars: int = 200) -> List[dict]:
    if not raw:
        return []
    result = []
    current = {"start": raw[0]["start"], "end": raw[0]["end"], "text": raw[0]["text"]}
    for seg in raw[1:]:
        gap = seg["start"] - current["end"]
        combined_len = len(current["text"]) + len(seg["text"])
        if gap <= max_gap and combined_len <= max_chars and len(current["text"]) < min_chars:
            current["text"] += seg["text"]
            current["end"] = seg["end"]
        else:
            result.append(current)
            current = {"start": seg["start"], "end": seg["end"], "text": seg["text"]}
    result.append(current)
    final = []
    for seg in result:
        if final and len(seg["text"]) < min_chars:
            prev = final[-1]
            gap = seg["start"] - prev["end"]
            if gap <= max_gap and len(prev["text"]) + len(seg["text"]) <= max_chars:
                prev["text"] += seg["text"]
                prev["end"] = seg["end"]
                continue
        final.append(seg)
    return final


def _clean(text: str) -> str:
    text = re.sub(r'<\|[^|]*\|>', '', text)
    text = re.sub(r'<\|/[^|]*\|>', '', text)
    return text.strip()


def _split_sentences(text: str) -> list:
    parts = re.split(r'(?<=[。！？；])', text)
    result = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if len(p) > 50:
            subs = re.split(r'(?<=[，,])', p)
            buf = ""
            for s in subs:
                s = s.strip()
                if not s:
                    continue
                if buf and len(buf) + len(s) > 50:
                    result.append(buf)
                    buf = s
                else:
                    buf = (buf + s) if buf else s
            if buf:
                result.append(buf)
        else:
            result.append(p)
    if not result:
        return [text]
    return result


def _classify_reason(text: str) -> str:
    t = text.lower()
    for kw in ['敏感', '脏话', '谩骂', '政治', '得罪', '不当言论', 'sensitive']:
        if kw in t: return "sensitive"
    for kw in ['离题', '无关', '打断', '准备', '调试', '寒暄', '休息', 'off_topic', 'offtopic']:
        if kw in t: return "off_topic"
    for kw in ['结巴', '重复', '自我重复', 'stutter']:
        if kw in t: return "stutter"
    for kw in ['冗余', '铺垫', '重复表述', 'redundant', '过渡']:
        if kw in t: return "redundant"
    for kw in ['附和', '嗯嗯', '对对', 'echo', '应和', '无观点', '催促']:
        if kw in t: return "echo"
    for kw in ['杂音', '笑声', '咳嗽', 'noise', '非语言']:
        if kw in t: return "noise"
    return "filler"


def _suggestion_to_isKept(suggestion: str) -> bool:
    # AI suggestions never execute deletions. Every segment starts as kept and
    # only a user's explicit checkbox decision may change isKept to False.
    return True


# ============================================
# Task persistence
# ============================================
def _save_task_to_disk(task_id: str):
    if task_id not in tasks:
        return
    task = tasks[task_id]
    if task.get("status") != "completed":
        return
    record = {
        "task_id": task_id,
        "filename": task.get("filename", ""),
        "created_at": task.get("created_at", ""),
        "file_path": task.get("file_path", ""),
        "segments": task.get("segments", []),
        "speakers": task.get("speakers", []),
        "chapters": task.get("chapters", []),
        "summary": task.get("summary", ""),
    }
    history = _load_history()
    found = False
    for i, h in enumerate(history):
        if h.get("task_id") == task_id:
            history[i] = record
            found = True
            break
    if not found:
        history.insert(0, record)
    try:
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Failed to save history: {e}")


def _load_history() -> List[dict]:
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return []


def _load_task_from_disk(task_id: str) -> Optional[dict]:
    history = _load_history()
    for h in history:
        if h.get("task_id") == task_id:
            return h
    return None


# ============================================
# Local ASR via FunASR
# ============================================
async def _local_transcribe(file_path: str, task_id: str) -> List[dict]:
    vad_model = await asyncio.to_thread(_get_local_vad)
    if vad_model is None or vad_model is False:
        raise Exception("VAD model not available")
    asr_model = await asyncio.to_thread(_get_local_asr)
    if asr_model is None:
        raise Exception("ASR model not available")

    logger.info(f"[{task_id}] Running VAD segmentation...")
    vad_result = await asyncio.wait_for(
        asyncio.to_thread(vad_model.generate, input=file_path),
        timeout=900.0,
    )

    from pydub import AudioSegment
    audio = AudioSegment.from_file(file_path)
    audio = audio.set_frame_rate(16000).set_channels(1).set_sample_width(2)

    speech_segments = []
    if vad_result and len(vad_result) > 0:
        for item in vad_result:
            segments = item.get("value", item.get("text", []))
            if isinstance(segments, list):
                for t in segments:
                    if isinstance(t, (list, tuple)) and len(t) >= 2:
                        start_ms, end_ms = int(t[0]), int(t[1])
                        dur_ms = end_ms - start_ms
                        if dur_ms < 500:
                            continue
                        speech_segments.append((start_ms, end_ms))

    if not speech_segments:
        logger.warning(f"[{task_id}] VAD found no speech, using whole file")
        speech_segments = [(0, len(audio))]

    logger.info(f"[{task_id}] VAD found {len(speech_segments)} speech segments")

    raw_segments = []
    total_segments = len(speech_segments)

    for seg_idx, (start_ms, end_ms) in enumerate(speech_segments):
        tasks[task_id]["progress"] = f"本地语音转文字... {seg_idx + 1}/{total_segments}"

        dur_ms = end_ms - start_ms
        if dur_ms < 300:
            continue

        chunk_audio = audio[start_ms:end_ms]
        if len(chunk_audio) < 100:
            continue

        chunk_path = file_path + f".chunk_{seg_idx}.wav"
        chunk_audio.export(chunk_path, format="wav")

        try:
            seg_result = await asyncio.wait_for(
                asyncio.to_thread(
                    asr_model.generate,
                    input=chunk_path,
                    language="zh",
                    use_itn=True,
                ),
                timeout=300.0,
            )

            if seg_result and len(seg_result) > 0:
                res = seg_result[0]
                if isinstance(res, dict):
                    seg_text = _clean(res.get("text", ""))
                    if seg_text and not seg_text.startswith("<|"):
                        sentences = _split_sentences(seg_text)
                        total = max(sum(len(s) for s in sentences), 1)
                        chunk_start = start_ms / 1000.0
                        chunk_end = end_ms / 1000.0
                        chunk_dur = chunk_end - chunk_start
                        t = chunk_start
                        for sentence in sentences:
                            s_dur = max((len(sentence) / total) * chunk_dur, 0.3)
                            raw_segments.append({
                                "start": round(t, 2),
                                "end": round(t + s_dur, 2),
                                "text": sentence,
                            })
                            t += s_dur
        except asyncio.TimeoutError:
            logger.warning(f"[{task_id}] Segment {seg_idx + 1}/{total_segments} timeout, skipped")
        except Exception as e:
            logger.warning(f"[{task_id}] Segment {seg_idx + 1}/{total_segments} error: {e}")
        finally:
            try:
                os.remove(chunk_path)
            except OSError:
                pass

    logger.info(f"[{task_id}] Local ASR raw segments: {len(raw_segments)}")
    return raw_segments


# ============================================
# DashScope ASR (云端转写，部署用)
# ============================================
async def _upload_to_dashscope_oss(file_path: str, filename: str) -> str:
    async with httpx.AsyncClient(timeout=httpx.Timeout(connect=120.0, read=600.0, write=1200.0, pool=120.0)) as http_client:
        policy_resp = await http_client.get(
            "https://dashscope.aliyuncs.com/api/v1/uploads",
            headers={"Authorization": f"Bearer {DASHSCOPE_API_KEY}"},
            params={"action": "getPolicy", "model": "sensevoice-v1"},
        )
        if policy_resp.status_code != 200:
            raise Exception(f"Failed to get upload policy: {policy_resp.status_code}")
        policy_data = policy_resp.json().get("data", {})
        upload_host = policy_data.get("upload_host", "")
        upload_dir = policy_data.get("upload_dir", "")
        if not upload_host or not upload_dir:
            raise Exception(f"Invalid upload policy: {policy_data}")
        key = f"{upload_dir}/{filename}"
        with open(file_path, "rb") as f:
            file_content = f.read()
        form_fields = {
            "OSSAccessKeyId": policy_data.get("oss_access_key_id", ""),
            "Signature": policy_data.get("signature", ""),
            "policy": policy_data.get("policy", ""),
            "x-oss-object-acl": policy_data.get("x_oss_object_acl", ""),
            "x-oss-forbid-overwrite": policy_data.get("x_oss_forbid_overwrite", ""),
            "key": key,
            "success_action_status": "200",
        }
        files = {"file": (filename, file_content)}
        upload_resp = await http_client.post(upload_host, data=form_fields, files=files)
        if upload_resp.status_code not in (200, 201, 204):
            raise Exception(f"OSS upload failed: {upload_resp.status_code}")
        return f"oss://{key}"


async def _dashscope_transcribe(file_path: str, filename: str, task_id: str) -> List[dict]:
    oss_url = await _upload_to_dashscope_oss(file_path, filename)
    async with httpx.AsyncClient(timeout=httpx.Timeout(connect=120.0, read=600.0, write=300.0, pool=120.0)) as http_client:
        submit_resp = await http_client.post(
            "https://dashscope.aliyuncs.com/api/v1/services/audio/asr/transcription",
            headers={
                "Authorization": f"Bearer {DASHSCOPE_API_KEY}",
                "Content-Type": "application/json",
                "X-DashScope-Async": "enable",
                "X-DashScope-OssResourceResolve": "enable",
            },
            json={
                "model": "sensevoice-v1",
                "input": {"file_urls": [oss_url]},
                "parameters": {"language_hints": ["zh"]},
            },
        )
        if submit_resp.status_code not in (200, 201):
            raise Exception(f"Transcription submit failed: {submit_resp.status_code}")
        submit_data = submit_resp.json()
        ds_task_id = submit_data.get("output", {}).get("task_id")
        if not ds_task_id:
            raise Exception(f"No task_id: {submit_data}")
        while True:
            poll_resp = await http_client.get(
                f"https://dashscope.aliyuncs.com/api/v1/tasks/{ds_task_id}",
                headers={"Authorization": f"Bearer {DASHSCOPE_API_KEY}"},
            )
            poll_data = poll_resp.json()
            status = poll_data.get("output", {}).get("task_status", "")
            if status == "SUCCEEDED":
                break
            elif status == "FAILED":
                raise Exception(f"Transcription failed: {poll_data.get('output', {}).get('message', '')}")
            await asyncio.sleep(3)

    raw_segments = []
    results = poll_data.get("output", {}).get("results", [])
    try:
        for r in results:
            url = r.get("transcription_url", "")
            if url:
                data = await _fetch_transcription_url(url)
                raw_segments = _parse_transcription_data(data)
                if raw_segments:
                    break
        if not raw_segments:
            for r in results:
                for t in r.get("transcripts", []):
                    text = t.get("text", "")
                    if not text:
                        continue
                    for s in t.get("sentences", []):
                        s_text = s.get("text", "")
                        if not s_text:
                            continue
                        raw_segments.append({
                            "start": round(s.get("begin_time", 0) / 1000.0, 2),
                            "end": round(s.get("end_time", 0) / 1000.0, 2),
                            "text": s_text.strip(),
                        })
    except Exception as e:
        logger.error(f"[{task_id}] Parse error: {e}")

    return raw_segments


async def _fetch_transcription_url(url: str) -> dict:
    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(connect=30.0, read=60.0, write=30.0, pool=30.0)) as c:
            r = await c.get(url)
            if r.status_code == 200:
                return r.json()
    except Exception:
        pass
    return {}


def _parse_transcription_data(data: dict) -> List[dict]:
    segments = []
    try:
        for t in data.get("transcripts", []):
            text = _clean(t.get("text", ""))
            if not text:
                continue
            for s in t.get("sentences", []):
                s_text = _clean(s.get("text", ""))
                if not s_text:
                    continue
                segments.append({
                    "start": round(s.get("begin_time", 0) / 1000.0, 2),
                    "end": round(s.get("end_time", 0) / 1000.0, 2),
                    "text": s_text,
                })
    except Exception:
        pass
    return segments


# ============================================
# Unified transcribe function
# ============================================
async def upload_and_transcribe(file_path: str, filename: str, task_id: str) -> List[dict]:
    raw_segments = []

    if ASR_MODE == "dashscope":
        logger.info(f"[{task_id}] Using DashScope ASR")
        tasks[task_id]["progress"] = "上传音频到云端... 5%"
        tasks[task_id]["percent"] = 5
        raw_segments = await _dashscope_transcribe(file_path, filename, task_id)
    else:
        logger.info(f"[{task_id}] Using LOCAL SenseVoice ASR")
        tasks[task_id]["progress"] = "本地语音转文字... 10%"
        tasks[task_id]["percent"] = 10
        raw_segments = await _local_transcribe(file_path, task_id)
        if not raw_segments:
            logger.error(f"[{task_id}] Local ASR returned empty results")
            return []

    merged = merge_segments(raw_segments, min_chars=20, max_gap=1.5, max_chars=50)
    segments = []
    for i, seg in enumerate(merged):
        segments.append({
            "id": i + 1, "speaker": "Speaker_A", "text": seg["text"],
            "startTime": seg["start"], "endTime": seg["end"],
            "suggestion": "keep", "isKept": True,
            "reason": None, "reasonDetail": None,
        })
    logger.info(f"[{task_id}] After merging: {len(segments)} segments")
    return segments


# ============================================
# DeepSeek analysis
# ============================================
def parse_json_response(response_text: str):
    for prefix in ["```json", "```"]:
        if response_text.startswith(prefix):
            response_text = response_text[len(prefix):]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    response_text = response_text.strip()

    obj_start = response_text.find('{')
    arr_start = response_text.find('[')
    if arr_start >= 0 and (obj_start < 0 or arr_start < obj_start):
        start = arr_start
        end = response_text.rfind(']')
    elif obj_start >= 0:
        start = obj_start
        end = response_text.rfind('}')
    else:
        return None

    if end > start:
        chunk = response_text[start:end + 1]
        try:
            return json.loads(chunk)
        except:
            pass
        try:
            return json.loads(re.sub(r',\s*([}\]])', r'\1', chunk))
        except:
            pass
    return None


async def _call_deepseek(prompt: str, task_id: str, label: str, retries: int = 3) -> Optional[str]:
    if deepseek_client is None:
        raise RuntimeError("缺少 DEEPSEEK_API_KEY 环境变量，无法执行云端文本分析。")
    for attempt in range(1, retries + 1):
        try:
            resp = await asyncio.wait_for(
                asyncio.to_thread(
                    deepseek_client.chat.completions.create,
                    model=LLM_MODEL,
                    messages=[{"role": "user", "content": prompt}],
                ),
                timeout=1200.0,
            )
            content = resp.choices[0].message.content
            if content and content.strip():
                return content.strip()
            logger.warning(f"[{task_id}] {label}: Model returned empty content (attempt {attempt})")
            if attempt < retries:
                await asyncio.sleep(3)
        except asyncio.TimeoutError:
            logger.warning(f"[{task_id}] {label}: Timeout (attempt {attempt}/{retries})")
            if attempt < retries:
                await asyncio.sleep(5)
        except Exception as e:
            err = str(e)
            logger.warning(f"[{task_id}] {label}: Error - {err[:200]} (attempt {attempt}/{retries})")
            if any(k in err.lower() for k in ['503', '429', 'overload', 'rate_limit', 'server_error']) and attempt < retries:
                await asyncio.sleep(10 * attempt)
            elif attempt >= retries:
                break
            else:
                await asyncio.sleep(5)
    return None


def _extract_speakers(segments: List[dict]) -> List[dict]:
    smap = {}
    for seg in segments:
        spk = seg.get("speaker", "Speaker_A")
        if spk not in smap:
            smap[spk] = {"id": spk, "role": seg.get("speakerRole", "guest"), "name": spk.replace("_", " ")}
        elif seg.get("speakerRole") == "host":
            smap[spk]["role"] = "host"
    return list(smap.values())


def _auto_chapter(segments: List[dict], gap: float = 30.0) -> List[dict]:
    if not segments:
        return []
    chapters = []
    ch_start = segments[0]["startTime"]
    ch_end = segments[0]["endTime"]
    ch_idx = 0
    for i in range(1, len(segments)):
        if segments[i]["startTime"] - segments[i - 1]["endTime"] >= gap:
            chapters.append({"title": segments[ch_idx]["text"][:15] + "...", "startTime": round(ch_start, 2), "endTime": round(ch_end, 2)})
            ch_start = segments[i]["startTime"]
            ch_idx = i
        ch_end = segments[i]["endTime"]
    chapters.append({"title": segments[ch_idx]["text"][:15] + "...", "startTime": round(ch_start, 2), "endTime": round(ch_end, 2)})
    if len(chapters) > 7:
        step = len(chapters) // 5
        chapters = chapters[::step][:7]
    return chapters


def _normalize_suggestion(val) -> str:
    s = str(val).lower().strip()
    if s in ("strong", "delete", "remove", "del", "删除", "强烈"):
        return "strong"
    if s in ("mild", "maybe", "consider", "一般", "建议", "可选"):
        return "mild"
    return "keep"


async def analyze_with_deepseek(segments: List[dict], task_id: str) -> dict:
    BATCH_SIZE = 25
    batches = []
    for batch_start in range(0, len(segments), BATCH_SIZE):
        batch = segments[batch_start:batch_start + BATCH_SIZE]
        batch_num = batch_start // BATCH_SIZE + 1
        total = (len(segments) + BATCH_SIZE - 1) // BATCH_SIZE
        lines = [f"[{format_time(s['startTime'])} - {format_time(s['endTime'])}] {s['speaker']}: {s['text']}" for s in batch]
        prompt = DEEPSEEK_BATCH_PROMPT.format(transcript_text="\n".join(lines))
        batches.append((batch_num, total, batch, prompt))

    sem = asyncio.Semaphore(3)
    results_map = {}

    async def process_batch(batch_num: int, total: int, batch: list, prompt: str):
        async with sem:
            tasks[task_id]["progress"] = f"DeepSeek 分析 Batch {batch_num}/{total}"
            response_text = await _call_deepseek(prompt, task_id, f"Batch {batch_num}/{total}")
            if not response_text:
                logger.warning(f"[{task_id}] Batch {batch_num}/{total}: returned empty, using defaults")
                return batch_num, list(batch)
            parsed = parse_json_response(response_text)
            if parsed and isinstance(parsed, dict):
                bs = parsed.get("segments", [])
                if bs and isinstance(bs, list):
                    if len(bs) != len(batch):
                        logger.warning(f"[{task_id}] Batch {batch_num}/{total}: expected {len(batch)} got {len(bs)}")
                    return batch_num, bs
                else:
                    logger.warning(f"[{task_id}] Batch {batch_num}/{total}: no segments key")
                    return batch_num, list(batch)
            elif parsed and isinstance(parsed, list):
                if len(parsed) != len(batch):
                    logger.warning(f"[{task_id}] Batch {batch_num}/{total}: expected {len(batch)} got {len(parsed)}")
                return batch_num, list(parsed)
            else:
                logger.warning(f"[{task_id}] Batch {batch_num}/{total}: parse failed. Raw: {response_text[:200]}")
                return batch_num, list(batch)

    coros = [process_batch(bn, bt, b, p) for bn, bt, b, p in batches]
    gathered = await asyncio.gather(*coros)
    for bn, result in sorted(gathered, key=lambda x: x[0]):
        results_map[bn] = result

    all_analyzed = []
    for bn in sorted(results_map):
        all_analyzed.extend(results_map[bn])

    result_segments = []
    for i, seg in enumerate(segments):
        if i < len(all_analyzed) and isinstance(all_analyzed[i], dict):
            a = all_analyzed[i]
            suggestion = _normalize_suggestion(a.get("suggestion", "keep"))
            if suggestion == "keep" and not a.get("isKept", True):
                suggestion = "strong"
            reason_val = a.get("reason") if suggestion != "keep" else None
            reason_detail = a.get("reasonDetail") if suggestion != "keep" else None
            if reason_val and reason_val not in ("filler", "echo", "noise", "redundant", "off_topic", "stutter", "sensitive"):
                reason_val = _classify_reason(str(reason_val))
            if suggestion != "keep" and not reason_val:
                reason_val = _classify_reason(reason_detail or "")
            result_segments.append({
                "id": seg["id"],
                "speaker": a.get("speaker", seg["speaker"]),
                "speakerRole": a.get("speakerRole", "guest"),
                "text": seg["text"],
                "startTime": seg["startTime"],
                "endTime": seg["endTime"],
                "suggestion": suggestion,
                "isKept": _suggestion_to_isKept(suggestion),
                "reason": reason_val,
                "reasonDetail": reason_detail or reason_val,
            })
        else:
            result_segments.append({**seg, "speakerRole": "guest", "reasonDetail": None})

    speakers = _extract_speakers(result_segments)

    valid_chapters = await _generate_chapters(result_segments, task_id)
    if not valid_chapters:
        valid_chapters = _auto_chapter(result_segments)

    summary = await _generate_summary(result_segments, task_id)

    return {"segments": result_segments, "speakers": speakers, "chapters": valid_chapters, "summary": summary}


async def _generate_chapters(segments: List[dict], task_id: str) -> List[dict]:
    if not segments:
        return []

    kept = [s for s in segments if s.get("isKept", True)]
    source = kept if len(kept) > 10 else segments

    lines = [f"[{format_time(s['startTime'])} - {format_time(s['endTime'])}] {s.get('speaker', 'Speaker_A')}: {s['text']}" for s in source]
    full_text = "\n".join(lines)

    if len(full_text) > 30000:
        step = max(1, len(source) // 200)
        sampled = source[::step]
        lines = [f"[{format_time(s['startTime'])} - {format_time(s['endTime'])}] {s.get('speaker', 'Speaker_A')}: {s['text']}" for s in sampled]
        full_text = "\n".join(lines)

    prompt = DEEPSEEK_CHAPTER_PROMPT.format(transcript_text=full_text)
    response_text = await _call_deepseek(prompt, task_id, "Chapter generation")

    if not response_text:
        return []

    parsed = parse_json_response(response_text)
    if not parsed:
        return []

    chapters = parsed.get("chapters", []) if isinstance(parsed, dict) else []
    if not chapters or not isinstance(chapters, list):
        return []

    valid = []
    for ch in chapters:
        if not isinstance(ch, dict) or not ch.get("title"):
            continue
        title = str(ch["title"]).strip()
        st = float(ch.get("startTime", 0))
        et = float(ch.get("endTime", 0))
        if et <= st:
            et = st + 60
        valid.append({"title": title, "startTime": round(st, 2), "endTime": round(et, 2)})

    if len(valid) > 7:
        valid = valid[:7]

    if valid:
        valid[0]["startTime"] = segments[0]["startTime"]
        valid[-1]["endTime"] = segments[-1]["endTime"]

    return valid


async def _generate_summary(segments: List[dict], task_id: str) -> str:
    if not segments:
        return ""

    kept = [s for s in segments if s.get("isKept", True)]
    source = kept if len(kept) > 10 else segments

    lines = [f"[{format_time(s['startTime'])} - {format_time(s['endTime'])}] {s.get('speaker', 'Speaker_A')}: {s['text']}" for s in source]
    full_text = "\n".join(lines)

    if len(full_text) > 30000:
        step = max(1, len(source) // 200)
        sampled = source[::step]
        lines = [f"[{format_time(s['startTime'])} - {format_time(s['endTime'])}] {s.get('speaker', 'Speaker_A')}: {s['text']}" for s in sampled]
        full_text = "\n".join(lines)

    prompt = DEEPSEEK_SUMMARY_PROMPT.format(transcript_text=full_text)
    response_text = await _call_deepseek(prompt, task_id, "Summary generation")

    if not response_text:
        return ""

    parsed = parse_json_response(response_text)
    if parsed and isinstance(parsed, dict):
        return str(parsed.get("summary", "")).strip()
    return ""


# ============================================
# Background task processor
# ============================================
async def _process_audio_task(task_id: str, file_path: str, filename: str):
    try:
        tasks[task_id]["status"] = "transcribing"
        tasks[task_id]["progress"] = "本地语音转文字..."
        segments = await upload_and_transcribe(file_path, filename, task_id)
        if not segments:
            raise Exception("语音转文字返回空结果，音频可能为静音或已损坏")
        tasks[task_id]["status"] = "analyzing"
        tasks[task_id]["progress"] = f"DeepSeek 文本分析... (共 {len(segments)} 段)"
        analysis = await analyze_with_deepseek(segments, task_id)
        tasks[task_id]["status"] = "completed"
        tasks[task_id]["progress"] = "完成!"
        tasks[task_id]["segments"] = analysis["segments"]
        tasks[task_id]["speakers"] = analysis["speakers"]
        tasks[task_id]["chapters"] = analysis["chapters"]
        tasks[task_id]["summary"] = analysis.get("summary", "")
        tasks[task_id]["message"] = f"处理完成，共 {len(analysis['segments'])} 段"
        _save_task_to_disk(task_id)
    except Exception as e:
        err_msg = str(e) or f"{type(e).__name__}: 未知错误"
        logger.error(f"[{task_id}] Failed: {err_msg}")
        logger.error(traceback.format_exc())
        tasks[task_id]["status"] = "failed"
        tasks[task_id]["error"] = err_msg
        tasks[task_id]["progress"] = f"处理失败: {err_msg}"


# ============================================
# Export helpers
# ============================================
def _generate_word(task_data: dict, speaker_names: dict) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    # Cross-platform CJK font: PingFang SC on macOS, Microsoft YaHei on Windows
    style.font.name = 'PingFang SC' if platform.system() == 'Darwin' else 'Microsoft YaHei'
    style.font.size = Pt(11)

    title = doc.add_heading('FX281 Studio 播客粗剪文稿', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    segments = task_data.get("segments", [])
    chapters = task_data.get("chapters", [])
    speakers = task_data.get("speakers", [])

    kept = [s for s in segments if s.get("isKept", True)]
    removed = [s for s in segments if not s.get("isKept", True)]
    mild_count = len([s for s in segments if s.get("suggestion") == "mild"])
    strong_count = len([s for s in segments if s.get("suggestion") == "strong"])

    p = doc.add_paragraph()
    summary = f'原始段数: {len(segments)} | 保留: {len(kept)} | 建议删减: {len(removed)}'
    if mild_count or strong_count:
        summary += f' (一般建议: {mild_count}, 强烈建议: {strong_count})'
    p.add_run(summary).font.size = Pt(9)
    speaker_lines = []
    for s in speakers:
        name = speaker_names.get(s["id"], s["id"])
        role = "主持人" if s["role"] == "host" else "被访人"
        speaker_lines.append(f"{name}({role})")
    p.add_run(f'\n说话人: {", ".join(speaker_lines)}').font.size = Pt(9)

    doc.add_paragraph()

    suggestion_labels = {"mild": "一般建议删减", "strong": "强烈建议删减"}

    if chapters:
        for i, ch in enumerate(chapters):
            doc.add_heading(f'{i+1}. {ch["title"]}', level=2)
            ch_segs = [s for s in segments if (s.get("startTime",0) >= ch.get("startTime",0) - 0.5) and (s.get("startTime",0) < ch.get("endTime",0) + 0.5)]
            for seg in ch_segs:
                p = doc.add_paragraph()
                name = speaker_names.get(seg.get("speaker",""), seg.get("speaker",""))
                time_str = format_time(seg.get("startTime", 0))
                run = p.add_run(f'[{time_str}] {name}: ')
                run.font.size = Pt(10)
                run.font.bold = True
                suggestion = seg.get("suggestion", "keep")
                if seg.get("isKept", True) and suggestion == "keep":
                    p.add_run(seg["text"]).font.size = Pt(10)
                elif suggestion == "mild":
                    run_text = p.add_run(seg["text"])
                    run_text.font.size = Pt(10)
                    run_text.font.color.rgb = RGBColor(128, 128, 128)
                    reason = seg.get("reason", "")
                    detail = seg.get("reasonDetail", "")
                    tag = suggestion_labels.get(suggestion, suggestion)
                    if reason or detail:
                        run_tag = p.add_run(f'  [{tag}] {reason}: {detail}')
                        run_tag.font.size = Pt(8)
                        run_tag.font.color.rgb = RGBColor(150, 150, 150)
                else:
                    run_del = p.add_run(seg["text"])
                    run_del.font.size = Pt(10)
                    run_del.font.strike = True
                    run_del.font.color.rgb = RGBColor(200, 50, 50)
                    reason = seg.get("reason", "")
                    detail = seg.get("reasonDetail", "")
                    tag = suggestion_labels.get(suggestion, suggestion)
                    if reason or detail:
                        run_tag = p.add_run(f'  [{tag}] {reason}: {detail}')
                        run_tag.font.size = Pt(8)
                        run_tag.font.color.rgb = RGBColor(150, 150, 150)
    else:
        for seg in segments:
            p = doc.add_paragraph()
            name = speaker_names.get(seg.get("speaker",""), seg.get("speaker",""))
            time_str = format_time(seg.get("startTime", 0))
            p.add_run(f'[{time_str}] {name}: ').font.size = Pt(10)
            suggestion = seg.get("suggestion", "keep")
            if seg.get("isKept", True) and suggestion == "keep":
                p.add_run(seg["text"]).font.size = Pt(10)
            elif suggestion == "mild":
                run_text = p.add_run(seg["text"])
                run_text.font.color.rgb = RGBColor(128, 128, 128)
            else:
                run_del = p.add_run(seg["text"])
                run_del.font.strike = True
                run_del.font.color.rgb = RGBColor(200, 50, 50)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _generate_mp3(task_data: dict, original_file_path: str) -> str:
    from pydub import AudioSegment
    audio = AudioSegment.from_file(original_file_path)
    segments = task_data.get("segments", [])
    kept_segments = [s for s in segments if s.get("isKept", True)]
    if not kept_segments:
        return original_file_path

    kept_segments.sort(key=lambda s: s.get("startTime", 0))
    result = AudioSegment.empty()
    for seg in kept_segments:
        start_ms = int(seg.get("startTime", 0) * 1000)
        end_ms = int(seg.get("endTime", 0) * 1000)
        if end_ms > start_ms:
            chunk = audio[start_ms:end_ms]
            result += chunk

    if len(result) == 0:
        return original_file_path

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    result.export(tmp.name, format="mp3", bitrate="192k")
    tmp.close()
    return tmp.name


# ============================================
# API Routes
# ============================================
@app.get("/")
async def root():
    return {
        "status": "ok", "service": "FX281 API", "version": "13.0.0",
        "asr_mode": ASR_MODE, "llm_provider": LLM_PROVIDER, "llm_model": LLM_MODEL,
    }


@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "asr_mode": ASR_MODE}


@app.get("/api/history")
async def get_history():
    history = _load_history()
    summaries = []
    for h in history:
        summaries.append({
            "task_id": h.get("task_id"),
            "filename": h.get("filename", ""),
            "created_at": h.get("created_at", ""),
            "segment_count": len(h.get("segments", [])),
            "kept_count": len([s for s in h.get("segments", []) if s.get("isKept", True)]),
            "mild_count": len([s for s in h.get("segments", []) if s.get("suggestion") == "mild"]),
            "strong_count": len([s for s in h.get("segments", []) if s.get("suggestion") == "strong"]),
        })
    return JSONResponse(content=summaries)


@app.get("/api/history/{task_id}")
async def get_history_detail(task_id: str):
    record = _load_task_from_disk(task_id)
    if not record:
        raise HTTPException(status_code=404, detail="Task not found in history")
    return JSONResponse(content=record)


@app.delete("/api/history/{task_id}")
async def delete_history(task_id: str):
    history = _load_history()
    target = None
    for h in history:
        if h.get("task_id") == task_id:
            target = h
            break
    if not target:
        raise HTTPException(status_code=404, detail="Task not found in history")

    new_history = [h for h in history if h.get("task_id") != task_id]
    try:
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(new_history, f, ensure_ascii=False, indent=2)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete: {e}")

    if task_id in tasks:
        file_path = tasks[task_id].get("file_path")
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
        del tasks[task_id]

    file_path = target.get("file_path", "")
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception:
            pass

    return JSONResponse(content={"status": "ok", "deleted": task_id})


@app.post("/api/process-audio")
async def process_audio(file: UploadFile = File(...)):
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded")

    allowed = ['.mp3', '.m4a', '.wav', '.flac', '.mp4', '.aac', '.ogg']
    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ''
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {ext}")

    tmp = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as f:
            content = await file.read()
            f.write(content)
            tmp = f.name
    except Exception as e:
        if tmp and os.path.exists(tmp): os.remove(tmp)
        raise HTTPException(status_code=500, detail=str(e))

    task_id = str(uuid.uuid4())
    tasks[task_id] = {
        "status": "uploaded", "progress": "文件已上传...",
        "segments": None, "speakers": None, "chapters": None,
        "summary": None, "error": None, "message": None,
        "filename": file.filename,
        "created_at": datetime.now().isoformat(),
        "file_path": tmp,
    }

    asyncio.create_task(_process_audio_task(task_id, tmp, file.filename))

    return JSONResponse(content={"task_id": task_id, "status": "uploaded"})


@app.get("/api/task/{task_id}")
async def get_task(task_id: str):
    task = tasks.get(task_id)
    if not task:
        record = _load_task_from_disk(task_id)
        if record:
            return JSONResponse(content={
                "task_id": task_id, "status": "completed",
                "segments": record.get("segments", []),
                "speakers": record.get("speakers", []),
                "chapters": record.get("chapters", []),
                "summary": record.get("summary", ""),
                "message": "从历史记录加载",
            })
        raise HTTPException(status_code=404, detail="Task not found")

    return JSONResponse(content={
        "task_id": task_id,
        "status": task.get("status"),
        "progress": task.get("progress"),
        "segments": task.get("segments"),
        "speakers": task.get("speakers"),
        "chapters": task.get("chapters"),
        "summary": task.get("summary"),
        "error": task.get("error"),
        "message": task.get("message"),
    })


@app.post("/api/task/{task_id}/decisions")
async def update_task_decisions(task_id: str, decision_payload: dict):
    """Persist explicit user keep/cut choices before preview or export."""
    task = tasks.get(task_id)
    loaded_from_history = False
    if not task:
        task = _load_task_from_disk(task_id)
        loaded_from_history = True
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    decisions = decision_payload.get("decisions", [])
    if not isinstance(decisions, list):
        raise HTTPException(status_code=400, detail="decisions must be an array")
    by_id = {
        str(item.get("id")): item.get("isKept")
        for item in decisions
        if isinstance(item, dict) and isinstance(item.get("isKept"), bool)
    }
    updated = 0
    for segment in task.get("segments", []):
        key = str(segment.get("id"))
        if key in by_id:
            segment["isKept"] = by_id[key]
            segment["humanDecision"] = "keep" if by_id[key] else "cut"
            updated += 1

    if loaded_from_history:
        history = _load_history()
        for index, record in enumerate(history):
            if record.get("task_id") == task_id:
                history[index] = task
                break
        with open(HISTORY_FILE, "w", encoding="utf-8") as file:
            json.dump(history, file, ensure_ascii=False, indent=2)
    else:
        _save_task_to_disk(task_id)

    return JSONResponse(content={"status": "ok", "updated": updated})


@app.post("/api/export/word/{task_id}")
async def export_word(task_id: str, speaker_names: str = "{}"):
    task = tasks.get(task_id)
    if not task or task.get("status") != "completed":
        record = _load_task_from_disk(task_id)
        if not record:
            raise HTTPException(status_code=404, detail="Task not found or not completed")
        task = record

    try:
        names = json.loads(speaker_names)
    except:
        names = {}

    try:
        file_path = _generate_word(task, names)
        return FileResponse(file_path, filename=f"FX281_{task_id[:8]}.docx",
                          media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export/mp3/{task_id}")
async def export_mp3(task_id: str):
    task = tasks.get(task_id)
    if not task or task.get("status") != "completed":
        record = _load_task_from_disk(task_id)
        if not record:
            raise HTTPException(status_code=404, detail="Task not found or not completed")
        task = record

    original_path = task.get("file_path", "")
    if not original_path or not os.path.exists(original_path):
        raise HTTPException(status_code=404, detail="Original audio file not found")

    try:
        file_path = _generate_mp3(task, original_path)
        return FileResponse(file_path, filename=f"FX281_edited_{task_id[:8]}.mp3",
                          media_type="audio/mpeg")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# 后端同源 serve 前端构建产物（dist/），部署时一个服务提供 API + 页面
from pathlib import Path as _Path
_dist_dir = _Path(__file__).parent.parent / "dist"
if _dist_dir.exists():
    app.mount("/", StaticFiles(directory=str(_dist_dir), html=True), name="frontend")
    logger.info(f"Serving frontend from: {_dist_dir}")


if __name__ == "__main__":
    import uvicorn
    logger.info("=" * 50)
    logger.info("FX281 Podcast Processing API v13.0")
    logger.info(f"ASR: {'Local SenseVoice (FunASR)' if ASR_MODE == 'local' else 'DashScope sensevoice-v1'}")
    logger.info(f"Analysis: {LLM_MODEL} ({LLM_PROVIDER})")
    logger.info(f"Data dir: {DATA_DIR}")
    logger.info("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))
